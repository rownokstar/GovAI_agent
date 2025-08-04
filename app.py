import streamlit as st
from langchain_core.messages import AIMessage, HumanMessage
from langchain_community.document_loaders import PyPDFLoader, CSVLoader, Docx2txtLoader, UnstructuredExcelLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
import os
from fpdf import FPDF
from docx import Document as DocxDocument
import pandas as pd
import io

# --- পেজ কনফিগারেশন ---
st.set_page_config(page_title="GovAI Pro 2.0", page_icon="🤖", layout="wide")
st.title("🤖 GovAI Pro 2.0 - আপনার ইন্টারেক্টিভ ডকুমেন্ট অ্যাসিস্ট্যান্ট")

# --- সাইডবার ---
with st.sidebar:
    st.header("🛠️ কনফিগারেশন")
    st.markdown("API Key-গুলো এখন Streamlit Secrets থেকে স্বয়ংক্রিয়ভাবে লোড হচ্ছে।")
    
    if st.button("💬 নতুন চ্যাট শুরু করুন", use_container_width=True):
        st.session_state.clear()
        st.rerun()

    st.markdown("---")
    st.info("এই অ্যাপটি বাংলা, ইংরেজি এবং বাংলিশ ভাষায় কাজ করতে পারে।")


# --- API কী লোড করা ---
groq_api_key = st.secrets.get("GROQ_API_KEY")
openai_api_key = st.secrets.get("OPENAI_API_KEY")

if not groq_api_key or not openai_api_key:
    st.error("অনুগ্রহ করে Streamlit Cloud-এর Settings > Secrets সেকশনে আপনার Groq এবং OpenAI API Key যোগ করুন।")
    st.stop()

# --- ফাংশনসমূহ ---

def get_vectorstore_from_file(uploaded_file):
    if "vector_store" in st.session_state and st.session_state.get("uploaded_file_name") == uploaded_file.name:
        return st.session_state.vector_store
    
    try:
        file_extension = os.path.splitext(uploaded_file.name)[1]
        temp_file_path = os.path.join("/tmp", uploaded_file.name)
        
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        if file_extension == ".pdf":
            loader = PyPDFLoader(temp_file_path)
        elif file_extension == ".docx":
            loader = Docx2txtLoader(temp_file_path)
        elif file_extension == ".csv":
            loader = CSVLoader(temp_file_path)
        elif file_extension in [".xls", ".xlsx"]:
            loader = UnstructuredExcelLoader(temp_file_path, mode="elements")
        else:
            st.error("এই ফাইল ফরম্যাটটি সাপোর্ট করে না।")
            return None
        
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        document_chunks = text_splitter.split_documents(documents)
        
        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        vector_store = FAISS.from_documents(document_chunks, embeddings)

        st.session_state.vector_store = vector_store
        st.session_state.uploaded_file_name = uploaded_file.name
        os.remove(temp_file_path)
        return vector_store
    except Exception as e:
        st.error(f"ফাইল প্রসেসিং-এ সমস্যা হয়েছে: {e}")
        return None

def get_context_retriever_chain(vector_store):
    llm = ChatGroq(temperature=0.2, groq_api_key=groq_api_key, model_name="Llama3-70b-8192")
    
    retriever = vector_store.as_retriever()
    
    prompt = ChatPromptTemplate.from_messages([
      ("system", "Answer the user's questions based on the below context:\n\n{context}"),
      ("user", "{input}"),
    ])
    
    stuff_documents_chain = create_stuff_documents_chain(llm, prompt)
    return create_retrieval_chain(retriever, stuff_documents_chain)

def get_conversational_rag_chain(retriever_chain):
    llm = ChatGroq(temperature=0.2, groq_api_key=groq_api_key, model_name="Llama3-70b-8192")
    
    prompt = ChatPromptTemplate.from_messages([
      ("system", """You are a helpful and friendly AI assistant named GovAI Pro. Answer the user's questions based on the provided context and conversation history.

Language Rules:
- If the user asks in English, respond in English.
- If the user asks in Bengali (বাংলা), respond in clear Bengali.
- If the user asks in Banglish (e.g., "amar kisu document lagbe"), understand it as Bengali and respond in clear Bengali.

Formatting Rules:
- Use appropriate emojis to make the conversation engaging. 👍
- If you need to present structured data (like comparisons, lists of items, etc.), use Markdown tables.
- If you don't know the answer from the context, politely say so in the correct language.

Context: {context}"""),
      ("user", "{input}"),
    ])
    
    stuff_documents_chain = create_stuff_documents_chain(llm, prompt)
    return create_retrieval_chain(retriever_chain.retriever, stuff_documents_chain)

# --- রিপোর্ট জেনারেট ফাংশন ---
def generate_pdf_report(chat_history):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
    pdf.set_font('DejaVu', '', 12)
    
    pdf.cell(200, 10, txt="GovAI Pro - Chat Report", ln=True, align='C')
    
    for message in chat_history:
        role = "User" if isinstance(message, HumanMessage) else "Assistant"
        # Correctly handle potential encoding issues for the content
        content = message.content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, f"{role}: {content}")
        pdf.ln(5)
        
    return pdf.output(dest='S').encode('latin-1')

def generate_docx_report(chat_history):
    document = DocxDocument()
    document.add_heading('GovAI Pro - Chat Report', 0)
    
    for message in chat_history:
        role = "User" if isinstance(message, HumanMessage) else "Assistant"
        document.add_paragraph(f"{role}: {message.content}")
        
    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()

def generate_excel_report(chat_history):
    data = {"Role": [], "Message": []}
    for message in chat_history:
        role = "User" if isinstance(message, HumanMessage) else "Assistant"
        data["Role"].append(role)
        data["Message"].append(message.content)
        
    df = pd.DataFrame(data)
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Chat History')
    return bio.getvalue()

# --- সেশন স্টেট শুরু করা ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = [AIMessage(content="नमस्ते! আমি GovAI Pro। আপনার ডকুমেন্ট আপলোড করুন এবং প্রশ্ন জিজ্ঞাসা শুরু করুন।")]

# --- ফাইল আপলোড এবং প্রসেসিং ---
uploaded_file = st.file_uploader(
    "📄 আপনার ডকুমেন্ট আপলোড করুন (PDF, DOCX, CSV, Excel)", 
    type=['pdf', 'docx', 'csv', 'xls', 'xlsx'],
    help="এখানে ফাইল আপলোড করার পর আপনি প্রশ্ন করতে পারবেন।"
)

if not uploaded_file:
    st.info("একটি ডকুমেন্ট আপলোড করে শুরু করুন।")
else:
    with st.spinner("আপনার ডকুমেন্ট প্রসেস করা হচ্ছে... ⏳"):
        vector_store = get_vectorstore_from_file(uploaded_file)
        if vector_store is not None:
            retriever_chain = get_context_retriever_chain(vector_store)

            # --- চ্যাট হিস্ট্রি ডিসপ্লে ---
            for message in st.session_state.chat_history:
                if isinstance(message, AIMessage):
                    with st.chat_message("AI", avatar="🤖"):
                        st.write(message.content)
                elif isinstance(message, HumanMessage):
                    with st.chat_message("Human", avatar="👤"):
                        st.write(message.content)

            # --- ইউজার ইনপুট ---
            user_query = st.chat_input("আপনার প্রশ্ন এখানে লিখুন...")
            if user_query:
                st.session_state.chat_history.append(HumanMessage(content=user_query))
                
                with st.chat_message("Human", avatar="👤"):
                    st.write(user_query)

                with st.chat_message("AI", avatar="🤖"):
                    # streaming response
                    def stream_response():
                        chain = get_conversational_rag_chain(retriever_chain)
                        response = chain.invoke({
                            "input": user_query,
                            "context": st.session_state.chat_history
                        })
                        return response['answer']

                    full_response = st.write_stream(stream_response)
                    st.session_state.chat_history.append(AIMessage(content=full_response))

                    # --- ডাউনলোড বাটন ---
                    st.markdown("---")
                    st.write("এই আলোচনাটি ডাউনলোড করুন:")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button(
                            label="📄 PDF হিসাবে",
                            data=generate_pdf_report(st.session_state.chat_history),
                            file_name="chat_report.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    with col2:
                        st.download_button(
                            label="📑 DOCX হিসাবে",
                            data=generate_docx_report(st.session_state.chat_history),
                            file_name="chat_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    with col3:
                        st.download_button(
                            label="📊 Excel হিসাবে",
                            data=generate_excel_report(st.session_state.chat_history),
                            file_name="chat_report.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
